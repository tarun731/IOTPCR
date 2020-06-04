from django.shortcuts import render
from .models import DCPUlogin, Organization, Branches, Children, Attendence, Problem, Funds, Problemsolution, Donations, Fundusagebills
from random import randint
from django.core.files.storage import FileSystemStorage
from datetime import date, datetime
import json
from django.http import JsonResponse
from django.core.files import File
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Create your views here.

def index(request):
    context = {
        'page' : 'index'
    }
    return render(request, 'govnportal/index.html', context)


def dcpulogin(request):
    link = 'govnportal/index.html'
    useriddata = request.POST.get('userid',False)
    password = request.POST.get('password',False)
    dcpulogindetails = DCPUlogin.objects.all()
    name = ''
    page = ''
    if useriddata == 'dcpu':
        name = 'District Child Protection Unit'
        page = 'dcpu'
    else:
        name = 'Children welfare commitee'
        page = 'cwc'

    for data in dcpulogindetails:
        print(data.userid)
        if data.userid == useriddata:
            if data.password == password:
                link = 'govnportal/dcpuloginpage.html'
                request.session['sessiondata'] = name
            else:
                link = 'govnportal/index.html'
    data = Organization.objects.all()
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'page1': page,
        'data': data,
        'page': 'home',
        'probs': no_active_probs
    }

    return render(request, link, context)

def dcpuhome(request):
    data = Organization.objects.all()
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'page': 'home',
        'data': data,
        'probs': no_active_probs
    }

    return render(request, 'govnportal/dcpuloginpage.html',context)


def viewbranches(request):
    orgid = request.POST.get('orgid', False)
    orgname = request.POST.get('orgname', False)
    print(orgid)
    orgobj = Organization.objects.get(orgid=orgid)
    data = Branches.objects.filter(organization=orgobj)
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'page': 'viewbranches',
        'data': data,
        'orgname': orgname,
        'probs': no_active_probs
    }
    return render(request, 'govnportal/dcpuloginpage.html', context)




def addorganization(request):
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'page': 'addorganization',
        'probs': no_active_probs
    }
    return render(request, 'govnportal/orgregister.html', context)

def orgindata(request):
    link = 'govnportal/orgregister.html'
    orgname = request.POST.get('orgname', False)
    ownername = request.POST.get('ownername', False)
    orgid = request.POST.get('orgid', False)
    password = request.POST.get('password', False)
    address = request.POST.get('address', False)
    village = request.POST.get('village', False)
    city = request.POST.get('city', False)
    state = request.POST.get('state', False)
    phoneno = request.POST.get('phoneno', False)
    orgidcheck = None
    id = ''
    status = ''
    try:
        orgidcheck = Organization.objects.get(orgid=orgid)
    except Organization.DoesNotExist:
        orgidcheck = None

    if orgidcheck == None:
        obj = Organization()
        obj.orgname = orgname
        obj.ownername = ownername
        obj.orgid = orgid
        obj.password = password
        obj.address = address
        obj.village = village
        obj.city = city
        obj.state = state
        obj.phoneno = phoneno
        obj.save()
        status = 'done'
    else:
        id = 'exists'

    data = Organization.objects.all()
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'page': id,
        'data': data,
        'status': status,
        'probs': no_active_probs
    }

    return render(request, link, context)

def dcpuviewattendance(request):
    link = 'govnportal/dcpuviewattendance.html'
    org_details = Organization.objects.all()
    att_check = 'yes'
    branch_check = 'yes'
    
    if request.method == "POST":
        body = request.body.decode('utf-8').replace('\0', '')
        x = json.loads(body)
        orgid = x['org']
        request.session['orgid_att'] = orgid
        print(orgid)
    
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'data_org': org_details, 
        'att_check': att_check,
        'branch_check': branch_check,
        'probs': no_active_probs
    }
    return render(request, link, context)

def dcpuviewatt_getbranch(request):
    link = 'govnportal/dcpuviewattendance.html'
    org_details = Organization.objects.all()
    orgid = request.session['orgid_att']
    print(orgid)
    orgobj = Organization.objects.get(orgid=orgid)
    branch_details = Branches.objects.filter(organization=orgobj)
    date = ''
    dataexists = ''
    att1 = []
    att_check = 'yes'
    branch_check = 'yes'
    if not branch_details:
        branch_check = 'no'
    if request.method == "POST":
        branchid = request.POST.get('branch',False)
        date = request.POST.get('date',False)
        branchobj = Branches.objects.get(branchid=branchid)
        att = Attendence.objects.filter(branch=branchobj)
        att1 = att.filter(date=date)
        if not att1:
            dataexists = 'no'
            att_check = 'no'
    
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'data_org': org_details,
        'data_branch': branch_details,
        'org': orgobj,
        'attendence': att1,
        'date': date,
        'dataexists': dataexists,
        'att_check': att_check,
        'branch_check': branch_check,
        'probs': no_active_probs
    }
    return render(request, link, context)

def viewprobs(request):
    link = 'govnportal/viewprobs.html'
    status = ''
    error = ''
    randomnum = randint(10000, 99999)
    problemid = request.POST.get('problem',False)
    problemobj = None
    todaydate = date.today()

    ids = None
    try:
        ids = Funds.objects.get(fundid=randomnum)
    except Funds.DoesNotExist:
        ids = None
        
    while ids!=None:
        randomnum = randint(10000, 99999)
        try:
            ids = Funds.objects.get(fundid=randomnum)
        except Funds.DoesNotExist:
            ids = None

    if request.method == "POST" and request.POST.get('orgname',False) == False:
        problem = request.POST.get('problem',False)
        problemobj = Problem.objects.get(problemid=problem)
        link = 'govnportal/dcpusolution.html'

    if request.method == "POST" and request.POST.get('orgname',False) != False:
        problemid = request.POST.get('problemid',False)
        problemname = request.POST.get('problemname',False)
        problemsol = request.POST.get('problemsolution',False)
        problemobj = Problem.objects.get(problemid=problemid)
        fundraise = request.POST.get('fundraise',False)
        print(fundraise)
    
        if fundraise == 'Yes':
            print('a')
            amount = request.POST.get('amount',False)
            paymentmode = request.POST.get('paymentmode',False)
            uploaded_file = request.FILES.get('photo',False)
            photo = uploaded_file.name
            if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
                error = 'no'
                fundid = randomnum
                print('b')

                fs = FileSystemStorage()
                fs.save('fundbill'+str(fundid)+'.jpg', uploaded_file)

                obj = Funds()
                obj.organization = problemobj.organization
                obj.branch = problemobj.branch
                obj.problem = problemobj
                obj.fundid = fundid
                obj.fundfor = problemname
                obj.fundraisedby = 'dcpu'
                obj.amount = amount
                obj.paymentmode = paymentmode
                obj.date = todaydate
                obj.balance = amount
                obj.completestatus = 'no'
                obj.paymentimage = 'fundbill'+str(fundid)+'.jpg'
                obj.save()

                fundobj = Funds.objects.get(fundid=fundid)

                objprob = Problemsolution()
                objprob.organization = problemobj.organization
                objprob.branch = problemobj.branch
                objprob.problem = problemobj
                objprob.fund = fundobj
                objprob.problemsolution = problemsol
                objprob.save()
                status = 'done'

                problemobj.solved = 'yes'
                problemobj.save()
            else:
                error = 'yes'
        else:
            objprob = Problemsolution()
            objprob.organization = problemobj.organization
            objprob.branch = problemobj.branch
            objprob.problem = problemobj
            objprob.problemsolution = problemsol
            objprob.save()
            status = 'done'

            problemobj.solved = 'yes'
            problemobj.save()
    

    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'data': active_probs1,
        'probs': no_active_probs,
        'data_problem': problemobj,
        'status': status,
        'error': error
    }
    return render(request, link, context)

def dcpusolvedprobs(request):
    link = 'govnportal/dcpusolvedprobs.html'
    data = Problem.objects.filter(orgacceptance='yes')
    data1 = data.filter(solved='yes')
    data1.reverse()

    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)

    solutions = []

    probsolutions = Problemsolution.objects.all()

    for a in data1:
        for b in probsolutions:
            if a == b.problem:
                solutions.append(b)

    solutions.reverse()

    context = {
        'data_probs': data1,
        'data_sols': solutions,
        'probs': no_active_probs,
    }
    return render(request, link, context)

def dcpuraisefund(request):
    link = 'govnportal/dcpuraisefund.html'
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    org_details = Organization.objects.all()
    branch_check = 'yes'
    if request.method == "POST":
        body = request.body.decode('utf-8').replace('\0', '')
        x = json.loads(body)
        orgid = x['org']
        request.session['orgid_att'] = orgid
    context = {
        'probs': no_active_probs,
        'data_org': org_details, 
        'branch_check': branch_check,
    }
    return render(request, link, context)


def dcpuraisefund_getbranch(request):
    link = 'govnportal/dcpuraisefund.html'
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    org_details = Organization.objects.all()
    orgid = request.session['orgid_att']
    orgobj = Organization.objects.get(orgid=orgid)
    branch_details = Branches.objects.filter(organization=orgobj)
    branch_check = 'yes'
    error = ''
    status = ''
    todaydate = date.today()
    if not branch_details:
        branch_check = 'no'
    if request.method == "POST":
        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Funds.objects.get(fundid=randomnum)
        except Funds.DoesNotExist:
            ids = None
        
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Funds.objects.get(fundid=randomnum)
            except Funds.DoesNotExist:
                ids = None

        uploaded_file = request.FILES['photo']
        photo = uploaded_file.name
        
        if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
            print('b')
            error = 'no'
            branchid = request.POST.get('branch',False)
            branchobj = Branches.objects.get(branchid=branchid)
            fundfor = request.POST.get('fundfor',False)
            fundid = randomnum
            fundamount = request.POST.get('amount',False)
            paymentmode = request.POST.get('paymentmode',False)

            fs = FileSystemStorage()
            fs.save('fundbill'+str(fundid)+'.jpg', uploaded_file)

            obj = Funds()
            obj.organization = orgobj
            obj.branch = branchobj
            obj.fundid = fundid
            obj.fundfor = fundfor
            obj.fundraisedby = 'dcpu'
            obj.amount = fundamount
            obj.paymentmode = paymentmode
            obj.date = todaydate
            obj.balance = fundamount
            obj.completestatus = 'no'
            obj.paymentimage = 'fundbill'+str(fundid)+'.jpg'
            obj.save()
            status = 'done'
        else:
            error = 'yes'
    
    context = {
        'org': orgobj,
        'probs': no_active_probs,
        'data_org': org_details, 
        'branch_check': branch_check,
        'data_branch': branch_details,
        'status': status,
        'error': error
    }
    return render(request, link, context)

def dcpuraisedfund(request):
    link = 'govnportal/dcpuraisedfund.html'
    funds = Funds.objects.filter(fundraisedby='dcpu')
    funds.reverse()
    prob_with_fund = Problemsolution.objects.all()
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    context = {
        'data_funds': funds,
        'probs': no_active_probs,
        'data_probsols': prob_with_fund
    }
    return render(request, link, context)


def dcpuviewbills(request):
    link = 'govnportal/dcpuviewbills.html'
    active_probs = Problem.objects.filter(solved='no')
    active_probs1 = active_probs.filter(orgacceptance='yes')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)
    bills = Fundusagebills.objects.all()
    bills.reverse()
    context = {
        'probs': no_active_probs,
        'data': bills
    }
    return render(request, link, context)



def logouttoindex(request):
    logoutpage = request.POST.get('logoutpage',False)
    if logoutpage == 'dcpu':
        del request.session['sessiondata']

    if logoutpage == 'org':
        del request.session['sessiondata']
        del request.session['sessiond']
        del request.session['orgid']

    if logoutpage == 'branch':
        del request.session['sessiondata']
        del request.session['sessiond']
        del request.session['branchid']

    context = {
        'page': 'index',
    }
    
    return render(request, 'govnportal/logout.html', context)


def orglogin(request):
    link = 'govnportal/index.html'
    useriddata = request.POST.get('userid',False)
    password = request.POST.get('password',False)
    orglogindetails = Organization.objects.all()
    i = 0
    orgid = None
    orgobj = None
    data = None
    datatype = ''
    no_active_probs = None

    for data in orglogindetails:
        if data.orgid == useriddata:
            if data.password == password:
                link = 'govnportal/orgloginpage.html'
                request.session['sessiond'] = 'orglogin'
                request.session['sessiondata'] = data.orgname
                request.session['orgid'] = data.orgid
                i = 1
            else:
                link = 'govnportal/index.html'
        else:
            link = 'govnportal/index.html'
    if i == 1:
        orgid = request.session['orgid']
        orgobj = Organization.objects.get(orgid=orgid)
        data = Branches.objects.filter(organization=orgobj)
        datatype = 'branches'
        active_probs = Problem.objects.filter(organization=orgobj)
        active_probs2 = active_probs.filter(orgacceptance='no')
        active_probs1 = active_probs2.filter(solved='no')
        no_active_probs = len(active_probs1)
    
    context = {
        'page': 'orglogin',
        'data': data,
        'datatype': datatype,
        'probs': no_active_probs

    }
    return render(request, link, context)


def orghome(request):
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    data = Branches.objects.filter(organization=orgobj)
    datatype = 'branches'
    if request.method == "POST":
        branchid = request.POST.get('branch',False)
        print(branchid)
        branchobj = Branches.objects.filter(branchid=branchid)
        data = Children.objects.filter(branch=branchobj)
        datatype = 'childrens'
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    context = {
        'data': data,
        'datatype': datatype,
        'probs': no_active_probs
    }
    return render(request, 'govnportal/orgloginpage.html',context)


def orgviewattendence(request):
    link = 'govnportal/orgviewattendence.html'
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    data = Branches.objects.filter(organization=orgobj)
    date = ''
    dataexists = ''
    att1 = []
    if request.method == "POST":
        branchid = request.POST.get('branch',False)
        print(branchid)
        date = request.POST.get('date',False)
        branchobj = Branches.objects.get(branchid=branchid)
        att = Attendence.objects.filter(branch=branchobj)
        att1 = att.filter(date=date)
        if not att1:
            dataexists = 'no'
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    context = {
        'data': data,
        'attendence': att1,
        'date': date,
        'dataexists': dataexists,
        'probs': no_active_probs
    }
    return render(request, link, context)

def addbranch(request):
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    context = {
        'page': 'addbranch',
        'probs': no_active_probs,

    }
    return render(request, 'govnportal/branchregister.html', context)


def branchindata(request):
    porgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=porgid)
    link = 'govnportal/Branchregister.html'
    status = ''
    branchname = request.POST.get('branchname', False)
    ownername = request.POST.get('ownername', False)
    branchid = request.POST.get('branchid', False)
    password = request.POST.get('password', False)
    address = request.POST.get('address', False)
    village = request.POST.get('village', False)
    city = request.POST.get('city', False)
    state = request.POST.get('state', False)
    phoneno = request.POST.get('phoneno', False)
    branchidcheck = None
    id = ''
    try:
        branchidcheck = Branches.objects.get(branchid=branchid)
    except Branches.DoesNotExist:
        branchidcheck = None
    

    if branchidcheck == None:
        organization = Organization.objects.get(orgid=porgid)
        obj = Branches()
        obj.organization = organization
        obj.branchname = branchname
        obj.ownername = ownername
        obj.branchid = branchid
        obj.password = password
        obj.address = address
        obj.village = village
        obj.city = city
        obj.state = state
        obj.phoneno = phoneno
        obj.save()
        status = 'done'
    else:
        id = 'exists'
    
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    context = {
        'page': id,
        'status': status,
        'probs': no_active_probs
    }

    return render(request, link, context)

def orgviewprobs(request):
    link = 'govnportal/orgviewprobs.html'
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    status = ''

    if request.method == "POST":
        problem = request.POST.get('problem',False)
        problemobj = Problem.objects.get(problemid=problem)
        problemobj.orgacceptance = 'yes'
        problemobj.save()
        status = 'done'

    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    active_probs1.reverse()
    context = {
        'data': active_probs1,
        'probs': no_active_probs,
        'status': status
    }
    return render(request, link, context)

def orgsolvedprobs(request):
    link = 'govnportal/orgsolvedprobs.html'
    porgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=porgid)
    data = Problem.objects.filter(organization=orgobj)
    data2 = data.filter(orgacceptance='no')
    data1 = data2.filter(solved='yes')
    data1.reverse()

    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(solved='no')
    active_probs1 = active_probs2.filter(orgacceptance='no')
    active_probs1.reverse()
    no_active_probs = len(active_probs1)

    solutions = []

    probsolutions = Problemsolution.objects.all()

    for a in data1:
        for b in probsolutions:
            if a == b.problem:
                solutions.append(b)

    solutions.reverse()

    context = {
        'data_probs': data1,
        'data_sols': solutions,
        'probs': no_active_probs,
    }
    return render(request, link, context)

def orgsolution(request):
    link = 'govnportal/orgsolution.html'
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    status = ''
    error = ''
    randomnum = randint(10000, 99999)
    problemid = request.POST.get('problem',False)
    todaydate = date.today()

    ids = None
    try:
        ids = Funds.objects.get(fundid=randomnum)
    except Funds.DoesNotExist:
        ids = None
        
    while ids!=None:
        randomnum = randint(10000, 99999)
        try:
            ids = Funds.objects.get(fundid=randomnum)
        except Funds.DoesNotExist:
            ids = None

    if request.method == "POST" and problemid == False:
        problemid = request.POST.get('problemid',False)
        problemname = request.POST.get('problemname',False)
        problemsol = request.POST.get('problemsolution',False)
        problemobj = Problem.objects.get(problemid=problemid)
        fundraise = request.POST.get('fundraise',False)
        print(fundraise)
    
        if fundraise == 'Yes':
            print('a')
            amount = request.POST.get('amount',False)
            paymentmode = request.POST.get('paymentmode',False)
            uploaded_file = request.FILES.get('photo',False)
            photo = uploaded_file.name
            if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
                error = 'no'
                fundid = randomnum
                print('b')

                fs = FileSystemStorage()
                fs.save('fundbill'+str(fundid)+'.jpg', uploaded_file)

                obj = Funds()
                obj.organization = orgobj
                obj.branch = problemobj.branch
                obj.problem = problemobj
                obj.fundid = fundid
                obj.fundfor = problemname
                obj.fundraisedby = 'org'
                obj.amount = amount
                obj.paymentmode = paymentmode
                obj.date = todaydate
                obj.balance = amount
                obj.completestatus = 'no'
                obj.paymentimage = 'fundbill'+str(fundid)+'.jpg'
                obj.save()

                fundobj = Funds.objects.get(fundid=fundid)

                objprob = Problemsolution()
                objprob.organization = orgobj
                objprob.branch = problemobj.branch
                objprob.problem = problemobj
                objprob.fund = fundobj
                objprob.problemsolution = problemsol
                objprob.save()
                status = 'done'

                problemobj.solved = 'yes'
                problemobj.save()
            else:
                error = 'yes'
        else:
            objprob = Problemsolution()
            objprob.organization = orgobj
            objprob.branch = problemobj.branch
            objprob.problem = problemobj
            objprob.problemsolution = problemsol
            objprob.save()
            status = 'done'

            problemobj.solved = 'yes'
            problemobj.save()
    else:
        problemobj = Problem.objects.get(problemid=problemid)
        

    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)

    context = {
        'data_problem': problemobj,
        'probs': no_active_probs,
        'status': status,
        'error': error,
    }
    return render(request, link, context)

def orgraisefund(request):
    link = 'govnportal/orgraisefund.html'
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    branches = Branches.objects.filter(organization=orgobj)
    todaydate = date.today()
    status = ''
    error = ''
    if request.method == "POST":
        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Funds.objects.get(fundid=randomnum)
        except Funds.DoesNotExist:
            ids = None
        
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Funds.objects.get(fundid=randomnum)
            except Funds.DoesNotExist:
                ids = None

        uploaded_file = request.FILES['photo']
        photo = uploaded_file.name
        
        if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
            error = 'no'
            branchid = request.POST.get('branch',False)
            branchobj = Branches.objects.get(branchid=branchid)
            fundfor = request.POST.get('fundfor',False)
            fundid = randomnum
            fundamount = request.POST.get('amount',False)
            paymentmode = request.POST.get('paymentmode',False)

            fs = FileSystemStorage()
            fs.save('fundbill'+str(fundid)+'.jpg', uploaded_file)

            obj = Funds()
            obj.organization = orgobj
            obj.branch = branchobj
            obj.fundid = fundid
            obj.fundfor = fundfor
            obj.fundraisedby = 'org'
            obj.amount = fundamount
            obj.date = todaydate
            obj.balance = fundamount
            obj.completestatus = 'no'
            obj.paymentmode = paymentmode
            obj.paymentimage = 'fundbill'+str(fundid)+'.jpg'
            obj.save()
            status = 'done'
        else:
            error = 'yes'

    context = {
        'data_branch': branches,
        'probs': no_active_probs,
        'status': status,
        'error': error
    }
    return render(request, link, context)

def orgraisedfund(request):
    link = 'govnportal/orgraisedfund.html'
    orgid = request.session['orgid']
    orgobj = Organization.objects.get(orgid=orgid)
    funds1 = Funds.objects.filter(fundraisedby='org')
    funds = funds1.filter(organization=orgobj)
    funds.reverse()
    prob_with_fund = Problemsolution.objects.all()
    active_probs = Problem.objects.filter(organization=orgobj)
    active_probs2 = active_probs.filter(orgacceptance='no')
    active_probs1 = active_probs2.filter(solved='no')
    no_active_probs = len(active_probs1)
    context = {
        'data_funds': funds,
        'probs': no_active_probs,
        'data_probsols': prob_with_fund
    }
    return render(request, link, context)


def branchlogin(request):
    link = 'govnportal/index.html'
    useriddata = request.POST.get('userid',False)
    password = request.POST.get('password',False)
    branchlogindetails = Branches.objects.all()
    i = 0
    data = None

    for data in branchlogindetails:
        if data.branchid == useriddata:
            if data.password == password:
                link = 'govnportal/branchloginpage.html'
                request.session['sessiond'] = 'branchlogin'
                request.session['sessiondata'] = data.branchname
                request.session['branchid'] = data.branchid
                i = 1
            else:
                link = 'govnportal/index.html'
        else:
            link = 'govnportal/index.html'
    if i == 1:
        branchid = request.session['branchid']
        branchobj = Branches.objects.get(branchid=branchid)
        data = Children.objects.filter(branch=branchobj)
    branchbal = branchobj.balance
    context = {
        'page': 'branchlogin',
        'data': data,
        'balance': branchbal
    }
    return render(request, link, context)


def branchhome(request):
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    data = Children.objects.filter(branch=branchobj)
    branchbal = branchobj.balance
    context = {
        'data': data,
        'balance': branchbal
    }
    return render(request, 'govnportal/branchloginpage.html',context)


def addchild(request):
    error = ''
    link = 'govnportal/childregister.html'
    status = ''
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    
    if request.method == 'POST':
        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Children.objects.get(childid=randomnum)
        except Children.DoesNotExist:
            ids = None
        
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Children.objects.get(childid=randomnum)
            except Children.DoesNotExist:
                ids = None

        uploaded_file = request.FILES['photo']
        photo = uploaded_file.name
        
        if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
            print('b')
            error = 'no'
            pbranchid = request.session['branchid']
            branch = Branches.objects.get(branchid=pbranchid)
            childid = randomnum
            childname = request.POST.get('childname',False)
            date = request.POST.get('date',False)
            gender = request.POST.get('gender',False)
            
            fs = FileSystemStorage()
            fs.save('cphoto'+str(childid)+'.jpg', uploaded_file)

            obj = Children()
            obj.branch = branch
            obj.childid = childid
            obj.childname = childname
            obj.childdob = date
            obj.gender = gender
            obj.photo = 'cphoto'+str(childid)+'.jpg'
            obj.save()
            status = 'complete'
        else:
            error = 'yes'

    branchbal = branchobj.balance
    context = {
        'error': error,
        'status': status,
        'balance': branchbal
    }
    return render(request, link, context)


def uploadattendence(request):
    link = 'govnportal/uploadattendence.html'
    status = ''
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    data = Children.objects.filter(branch=branchobj)
    todaydate = date.today()
    att_details = Attendence.objects.filter(branch=branchobj)
    att_details_today = att_details.filter(date=todaydate)
    att_name = "Upload Morning Attendence"
    att_type = request.POST.get('night',False)
    att_check = ""
    existids = []
    pen_visit_name = ""

    for a in att_details_today:
        existids.append(a.child)

    if len(att_details_today) == len(data):
        i = 0
        print(i)
        for a in att_details_today:
            if a.morningattendence == None:
                att_name = "Upload Morning Attendence"
            elif a.guestintime != None and a.guestouttime == None:
                att_name = "pending"
                i += 1
                pen_visit_name = "Guest visit"
            elif a.pvouttime != None and a.pvintime == None:
                att_name = "pending"
                i += 1
                pen_visit_name = "Personal visit"
            elif a.schoolouttime != None and a.schoolintime == None:
                att_name = "pending"
                i += 1
                pen_visit_name = "School visit"
            elif i == 0 and a.nightattendence == None:
                att_name = "Upload Night Attendence"
                print('a')
            elif a.nightattendence != None:
                att_check = "already Uploaded"
        
    elif len(att_details_today) < len(data):
        att_name = "Upload Morning Attendence"

    if request.method == "POST" and att_type == False:
        print('a')

        if not att_details_today:
            print('aa')
            for a in data:
                obj = Attendence()
                obj.branch = branchobj
                obj.child = Children.objects.get(childid=a.childid)
                obj.date = todaydate
                atted = str(a.childid)
                obj.morningattendence = request.POST.get(atted,False)
                obj.save()
                
        else:
            print('aaa')
            for a in data:
                if a not in existids:
                    print('abc')
                    obj = Attendence()
                    obj.branch = branchobj
                    childobj = Children.objects.get(childid=a.childid)
                    obj.child = childobj
                    obj.date = todaydate
                    atted = str(a.childid)
                    obj.morningattendence = request.POST.get(atted,False)
                    obj.save()
                    existids.append(a)
            for b in att_details_today:
                print('def')
                atted = str(a.childid)
                b.morningattendence = request.POST.get(atted,False)
                b.save()
        status = 'uploaded'


    elif request.method == "POST" and att_type == 'night':
        for a in data:
            atted = str(a.childid)
            childobj = Children.objects.filter(childid=a.childid)
            upp_data = Attendence.objects.filter(child=childobj)
            upp_data1 = upp_data.get(date=todaydate)
            upp_data1.nightattendence = request.POST.get(atted,False)
            upp_data1.save()
        status = 'uploaded'
    branchbal = branchobj.balance
    context = {
        'data': data,
        'att_name': att_name,
        'att_check': att_check,
        'status': status,
        'penvisit': pen_visit_name,
        'balance': branchbal
    }

    return render(request, link, context)


def guestvisit(request):
    link = 'govnportal/guestvisit.html'
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    data = Children.objects.filter(branch=branchobj)
    todaydate = date.today()
    time = datetime.now().strftime('%H:%M:%S')
    guest_details = Attendence.objects.filter(branch=branchobj)
    guest_details1 = guest_details.filter(date=todaydate)
    active_guests = []
    nonactive_guests = []
    attendence = ""
    status = ""

    if request.method == "POST":
        guestname = request.POST.get('guestname',False)
        mobileno = request.POST.get('mobileno',False)
        child = request.POST.get('child',False)
        childobjs = Children.objects.get(childid=child)
        print(childobjs)
        if not guest_details1:
            obj = Attendence()
            obj.branch = branchobj
            obj.child = childobjs
            obj.date = todaydate
            obj.guestname = guestname
            obj.guestmno = mobileno
            obj.guestintime = time
            obj.save()
            status = "uploaded"
            
        else:
            att_obj = guest_details1.get(child=childobjs)
            att_obj.guestname = guestname
            att_obj.guestmno = mobileno
            att_obj.guestintime = time
            att_obj.save()
            status = "uploaded"
    
    guest_details_aff = Attendence.objects.filter(branch=branchobj)
    guest_details_aff1 = guest_details_aff.filter(date=todaydate)

    if guest_details_aff1:
        for a in guest_details_aff1:
            if a.nightattendence != None:
                attendence = "complete"
        for a in guest_details_aff1:
            if a.guestname != None and a.guestintime != None and a.guestouttime == None:
                active_guests.append(a)

        for a in guest_details_aff1:
            if a.guestname == None and a.guestintime == None and a.guestouttime == None:
                nonactive_guests.append(a.child)
    else:
        nonactive_guests = data
    branchbal = branchobj.balance
    context = {
        'data': nonactive_guests,
        'activeguests': active_guests,
        'att': attendence,
        'status': status,
        'balance': branchbal
    }
    return render(request, link, context)


def delguest(request):
    guest = request.POST.get('guest')
    childobj = Children.objects.get(childid=guest)
    time = datetime.now().strftime('%H:%M:%S')
    todaydate = date.today()
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    guest_details = Attendence.objects.filter(branch=branchobj)
    guest_details1 = guest_details.filter(date=todaydate)
    for a in guest_details1:
        if a.child == childobj:
            print('a')
            a.guestouttime = time
            a.save()
    status = 'deleted'
    active_guests = []
    nonactive_guests = []
    attendence = ""
    guest_details_aff = Attendence.objects.filter(branch=branchobj)
    guest_details_aff1 = guest_details_aff.filter(date=todaydate)
    data = Children.objects.filter(branch=branchobj)

    if guest_details_aff1:
        for a in guest_details_aff1:
            if a.nightattendence != None:
                attendence = "complete"
        for a in guest_details_aff1:
            if a.guestname != None and a.guestintime != None and a.guestouttime == None:
                active_guests.append(a)

        for a in guest_details_aff1:
            if a.guestname == None and a.guestintime == None and a.guestouttime == None:
                nonactive_guests.append(a.child)
    else:
        nonactive_guests = data
    branchbal = branchobj.balance
    context = {
        'data': nonactive_guests,
        'activeguests': active_guests,
        'att': attendence,
        'status': status,
        'balance': branchbal
    }
    return render(request,'govnportal/guestvisit.html',context)


def personalvisit(request):
    link = 'govnportal/personalvisit.html'
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    data = Children.objects.filter(branch=branchobj)
    todaydate = date.today()
    time = datetime.now().strftime('%H:%M:%S')
    pers_details = Attendence.objects.filter(branch=branchobj)
    pers_details1 = pers_details.filter(date=todaydate)
    active_per_childs = []
    nonactive_per_childs = []
    attendence = ""
    status = ''
        
    if request.method == "POST":
        purpose = request.POST.get('visitpurpose',False)
        child = request.POST.get('child',False)
        childobjs = Children.objects.get(childid=child)
        if pers_details1:
            att_obj = pers_details1.get(child=childobjs)
            print(att_obj)
            if not att_obj:
                obj = Attendence()
                obj.branch = branchobj
                obj.child = childobjs
                obj.date = todaydate
                obj.personalvisit = purpose
                obj.pvouttime = time
                obj.save()
                status = 'uploaded'
                
            else:
                att_obj.personalvisit = purpose
                att_obj.pvouttime = time
                att_obj.save()
                status = 'uploaded'
        else:
            obj = Attendence()
            obj.branch = branchobj
            obj.child = childobjs
            obj.date = todaydate
            obj.personalvisit = purpose
            obj.pvouttime = time
            obj.save()
            status = 'uploaded'

    pers_details = Attendence.objects.filter(branch=branchobj)
    pers_details1 = pers_details.filter(date=todaydate)

    if pers_details1:
        for a in pers_details1:
            if a.nightattendence != None:
                attendence = "complete"
        for a in pers_details1:
            if a.personalvisit != None and a.pvouttime != None and a.pvintime == None:
                active_per_childs.append(a)
        if len(pers_details1) != len(data):
            nonactive_per_childs = data 
        else:
            for a in pers_details1:
                if a.personalvisit == None and a.pvouttime == None and a.pvintime == None:
                    nonactive_per_childs.append(a.child)
    else:
        nonactive_per_childs = data

    branchbal = branchobj.balance
    context = {
        'data': data,
        'active_per_childs': active_per_childs,
        'nonactive_per_childs': nonactive_per_childs,
        'att': attendence,
        'status': status,
        'balance': branchbal
    }
    return render(request,link,context)


def delvisit(request):
    child = request.POST.get('child')
    childobj = Children.objects.get(childid=child)
    time = datetime.now().strftime('%H:%M:%S')
    todaydate = date.today()
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    pers_details = Attendence.objects.filter(branch=branchobj)
    pers_details1 = pers_details.filter(date=todaydate)
    for a in pers_details1:
        if a.child == childobj:
            print('a')
            a.pvintime = time
            a.save()
    status = 'deleted'
    active_per_childs = []
    nonactive_per_childs = []
    attendence = ""
    data = Children.objects.filter(branch=branchobj)
    pers_details = Attendence.objects.filter(branch=branchobj)
    pers_details1 = pers_details.filter(date=todaydate)

    if pers_details1:
        for a in pers_details1:
            if a.nightattendence != None:
                attendence = "complete"
        for a in pers_details1:
            if a.personalvisit != None and a.pvouttime != None and a.pvintime == None:
                active_per_childs.append(a)
        if len(pers_details1) != len(data):
            nonactive_per_childs = data
        else:
            for a in pers_details1:
                if a.personalvisit == None and a.pvouttime == None and a.pvintime == None:
                    nonactive_per_childs.append(a.child)
    branchbal = branchobj.balance
    context = {
        'data': data,
        'active_per_childs': active_per_childs,
        'nonactive_per_childs': nonactive_per_childs,
        'att': attendence,
        'status': status,
        'balance': branchbal
    }
    return render(request,'govnportal/personalvisit.html',context)


def viewattendence(request):
    link = 'govnportal/viewattendence.html'
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    todaydate = date.today()
    att = Attendence.objects.filter(branch=branchobj)
    att1 = att.filter(date=todaydate)
    if request.method == "POST":
        todaydate = request.POST.get('date')
        att1 = att.filter(date=todaydate)
    branchbal = branchobj.balance
    context = {
        'data': att1,
        'date': todaydate,
        'balance': branchbal
    }
    return render(request,link,context)

def schoolvisit(request):
    link = 'govnportal/schoolvisit.html'
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    time = datetime.now().strftime('%H:%M:%S')
    todaydate = date.today()
    page = ""
    att_check = ""
    att = Attendence.objects.filter(branch=branchobj)
    att1 = att.filter(date=todaydate)
    active_visits = []
    nonactive_visits = []
    status = ''
    attendence = ""
    if request.method == "POST":
        if att1:
            i = 0
            for a in att1:
                atted = str(a.child.childid)
                child = request.POST.get(atted, False)
                print(child)
                if child != False:
                    i += 1
                    print(i)
                    a.schoolouttime = time
                    a.save()
                    status = 'done'
            if i == 0:
                page = "not selected"
    att = Attendence.objects.filter(branch=branchobj)
    att1 = att.filter(date=todaydate)
    if not att1:
        att_check = "notuploaded"
    elif att1:
        for a in att1:
            if a.nightattendence != None:
                attendence = "complete"
            if a.schoolouttime != None and a.schoolintime == None:
                active_visits.append(a.child)
            if a.schoolouttime == None and a.schoolintime == None:
                nonactive_visits.append(a.child)
    branchbal = branchobj.balance
    context = {
        'data': nonactive_visits,
        'page': page,
        'att_check': att_check,
        'active_visits': active_visits,
        'status': status,
        'attendence': attendence,
        'balance': branchbal
    }
    return render(request,link,context)


def schoolvisitcomplete(request):
    link = 'govnportal/schoolvisit.html'
    child = request.POST.get('child',False)
    print(child)
    childobj = Children.objects.get(childid=child)
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    time = datetime.now().strftime('%H:%M:%S')
    todaydate = date.today()
    att = Attendence.objects.filter(branch=branchobj)
    att1 = att.filter(date=todaydate)
    att2 = att1.get(child=childobj)
    att2.schoolintime = time
    att2.save()
    status = 'complete'
    att_check = ""
    active_visits = []
    nonactive_visits = []
    att = Attendence.objects.filter(branch=branchobj)
    att1 = att.filter(date=todaydate)
    if not att1:
        att_check = "notuploaded"
    elif att1:
        for a in att1:
            if a.schoolouttime != None and a.schoolintime == None:
                active_visits.append(a.child)
            if a.schoolouttime == None and a.schoolintime == None:
                nonactive_visits.append(a.child)
    branchbal = branchobj.balance
    context = {
        'data': nonactive_visits,
        'att_check': att_check,
        'active_visits': active_visits,
        'status': status,
        'balance': branchbal
    }
    return render(request,link,context)

def uploadproblem(request):
    link = 'govnportal/uploadproblem.html'
    status = ''
    branchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=branchid)
    todaydate = date.today()
    if request.method == "POST":
        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Problem.objects.get(problemid=randomnum)
        except Problem.DoesNotExist:
            ids = None
        
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Problem.objects.get(problemid=randomnum)
            except Problem.DoesNotExist:
                ids = None

        uploaded_file = request.FILES['photo']
        photo = uploaded_file.name
        
        if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
            pbranchid = request.session['branchid']
            branchobj = Branches.objects.get(branchid=pbranchid)
            problemid = randomnum
            probname = request.POST.get('probname',False)
            probdes = request.POST.get('probdescription',False)
            solved = 'no'
            
            fs = FileSystemStorage()
            fs.save('prob'+str(problemid)+'.jpg', uploaded_file)

            obj = Problem()
            obj.organization = branchobj.organization
            obj.branch = branchobj
            obj.problemname = probname
            obj.problemdes = probdes
            obj.problemid = problemid
            obj.photo = 'prob'+str(problemid)+'.jpg'
            obj.date = todaydate
            obj.orgacceptance = 'no'
            obj.solved = solved
            obj.save()
            status = 'done'
    branchbal = branchobj.balance
    context = {
        'status': status,
        'balance': branchbal
    }
    return render(request,link,context)

def unsolvedprob(request):
    link = 'govnportal/unsolvedprob.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    problems = Problem.objects.filter(branch=branchobj)
    problems1 = problems.filter(solved='no')
    branchbal = branchobj.balance
    context = {
        'data': problems1,
        'balance': branchbal
    }
    return render(request,link,context)

def solvedprob(request):
    link = 'govnportal/solvedprob.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    problems = Problemsolution.objects.filter(branch=branchobj)
    branchbal = branchobj.balance
    context = {
        'data': problems,
        'balance': branchbal
    }
    return render(request,link,context)

def donations(request):
    link = 'govnportal/donations.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    status = ''
    data = None

    if request.method == "POST":
        donorname = request.POST.get('donorname',False)
        phnno = request.POST.get('phn_no',False)
        address = request.POST.get('address',False)
        donationform = request.POST.get('type',False)
        amount = request.POST.get('amount',False)
        foodtime = request.POST.get('time')
        todaydate = date.today()
        strdate = str(todaydate)
        print(strdate)

        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Donations.objects.get(donationid=randomnum)
        except Donations.DoesNotExist:
            ids = None
            
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Donations.objects.get(donationid=randomnum)
            except Donations.DoesNotExist:
                ids = None

        donationid = randomnum

        document = docx.Document()

        document.add_heading('RECEIPT', 0).bold = True


        p_date1 = document.add_paragraph()
        p_date1.add_run('Date : ').bold = True
        p_date1.add_run(strdate)

        p_name = document.add_paragraph()
        p_name.add_run('Donor Name : ').bold = True
        p_name.add_run(donorname)

        p_phn = document.add_paragraph()
        p_phn.add_run('Donor Mobile NO : ').bold = True
        p_phn.add_run(phnno)

        p_add = document.add_paragraph()
        p_add.add_run('Address : ').bold = True
        p_add.add_run(address)

        p_type = document.add_paragraph()
        p_type.add_run('Donation Type : ').bold = True
        p_type.add_run(donationform)

        if donationform == 'money':
            p_amount = document.add_paragraph()
            p_amount.add_run('Donated Amount : ').bold = True
            p_amount.add_run(amount)
        elif donationform == 'food':
            p_food = document.add_paragraph()
            p_food.add_run('Food Prepared for : ').bold = True
            p_food.add_run(foodtime)
        
        breakline = document.add_paragraph()
        run = breakline.add_run()
        run.add_break()

        p_3 = document.add_paragraph()
        p_3.add_run('Admin Signature').bold = True
        p_3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        p_1 = document.add_paragraph()
        p_1.add_run('Regards     ').bold = True
        p_1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        p_2 = document.add_paragraph()
        p_2.add_run(branchobj.branchname).bold = True
        p_2.add_run('    ').bold = True
        p_2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        document.save('media/receipt_'+str(donationid)+'.docx')

        obj = Donations()
        obj.organization = branchobj.organization
        obj.branch = branchobj
        obj.donationid = donationid
        obj.donorname = donorname
        obj.phoneno = phnno
        obj.address = address
        obj.informof = donationform
        if donationform == 'money':
            obj.amount = amount
        elif donationform == 'food':
            obj.foodtime = foodtime
        obj.date = todaydate
        obj.receipt = 'receipt_'+str(donationid)+'.docx'
        obj.save()

        if donationform == 'money':
            if branchobj.balance == None:
                branchobj.balance = amount
            else:
                branchobj.balance = branchobj.balance + int(amount)
            branchobj.save()

        status = 'done'

        data = Donations.objects.get(donationid=donationid)

    branchbal = branchobj.balance
    context = {
        'status': status,
        'data': data,
        'balance': branchbal
    }
    return render(request,link,context)

def viewdonations(request):
    link = 'govnportal/viewdonations.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)

    donations = Donations.objects.filter(branch=branchobj)
    branchbal = branchobj.balance
    context = {
        'data': donations,
        'balance': branchbal
    }
    return render(request,link,context)

def branchraisedfunds(request):
    link = 'govnportal/branchraisedfunds.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    funds = Funds.objects.filter(branch=branchobj)
    branchbal = branchobj.balance
    context = {
        'data': funds,
        'balance': branchbal
    }
    return render(request,link,context)

def spentbills(request):
    link = 'govnportal/spentbills.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    funds1 = Funds.objects.filter(branch=branchobj)
    funds = funds1.filter(completestatus='no')
    status = ''
    error = ''
    error1 = ''
    error2 = ''
    todaydate = date.today()

    if request.method == "POST":
        fundid = request.POST.get('fund',False)

        if fundid != 'abc':
            amount = request.POST.get('amount',False)
            amount = int(amount)
            status = request.POST.get('status',False)
            fundobj = Funds.objects.get(fundid=fundid)

            randomnum = randint(10000, 99999)

            ids = None
            try:
                ids = Fundusagebills.objects.get(bill='spentbill'+str(randomnum)+'.jpg')
            except Fundusagebills.DoesNotExist:
                ids = None
            
            while ids!=None:
                randomnum = randint(10000, 99999)
                try:
                    ids = Fundusagebills.objects.get(bill='spentbill'+str(randomnum)+'.jpg')
                except Fundusagebills.DoesNotExist:
                    ids = None

            uploaded_file = request.FILES['bill']
            photo = uploaded_file.name
            
            if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
                if amount <= fundobj.balance or amount <= fundobj.balance + branchobj.balance:
                    error = 'no'
                    fs = FileSystemStorage()
                    fs.save('spentbill'+str(randomnum)+'.jpg', uploaded_file)

                    obj = Fundusagebills()
                    obj.organization = branchobj.organization
                    obj.branch = branchobj
                    obj.fund = fundobj
                    obj.bill = 'spentbill'+str(randomnum)+'.jpg'
                    obj.amount = amount
                    obj.date = todaydate
                    obj.save()

                    if amount <= fundobj.balance:
                        fundobj.balance = fundobj.balance - amount
                        fundobj.save()
                        if fundobj.balance == 0:
                            fundobj.completestatus = 'yes'
                            fundobj.save()
                        if status == 'yes':
                            fundobj.completestatus = 'yes'
                            fundobj.save()
                            if branchobj.balance == None:
                                branchobj.balance = fundobj.balance
                            else:
                                branchobj.balance = branchobj.balance + fundobj.balance
                            branchobj.save()
                            fundobj.balance = 0
                            fundobj.save()
                    elif amount <= fundobj.balance + branchobj.balance:
                        outstandingamount = amount - fundobj.balance
                        fundobj.balance = 0
                        fundobj.completestatus = 'yes'
                        fundobj.save()
                        branchobj.balance = branchobj.balance - outstandingamount
                        branchobj.save()

                    status = 'done'
                        
                else:
                    error2 = 'yes'
            else:
                error = 'yes'
        else:
            error1 = 'yes'
    branchbal = branchobj.balance
    context = {
        'data': funds,
        'status': status,
        'error': error,
        'error1': error1,
        'error2': error2,
        'balance': branchbal
    }
    return render(request,link,context)

def uploadedbills(request):
    link = 'govnportal/uploadedbills.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    bills = []

    fundbills = Fundusagebills.objects.all()

    for a in fundbills:
        if a.fund:
            if a.branch == branchobj:
                bills.append(a)

    bills.reverse()
    branchbal = branchobj.balance
    context = {
        'data': bills,
        'balance': branchbal
    }
    return render(request,link,context)

def uploadbills(request):
    link = 'govnportal/uploadbills.html'
    pbranchid = request.session['branchid']
    branchobj = Branches.objects.get(branchid=pbranchid)
    todaydate = date.today()
    error = ''
    status = ''
    if request.method == "POST":
        amount = request.POST.get('amount', False)
        amount = int(amount)

        randomnum = randint(10000, 99999)

        ids = None
        try:
            ids = Fundusagebills.objects.get(bill='spentbill'+str(randomnum)+'.jpg')
        except Fundusagebills.DoesNotExist:
            ids = None
            
        while ids!=None:
            randomnum = randint(10000, 99999)
            try:
                ids = Fundusagebills.objects.get(bill='spentbill'+str(randomnum)+'.jpg')
            except Fundusagebills.DoesNotExist:
                ids = None

        uploaded_file = request.FILES['bill']
        photo = uploaded_file.name

        if '.JPG' in photo or '.png' in photo or '.jpeg' in photo or '.jpg' in photo or '.PNG' in photo or '.JPEG' in photo:
            error = 'no'
            fs = FileSystemStorage()
            fs.save('bill'+str(randomnum)+'.jpg', uploaded_file)

            obj = Fundusagebills()
            obj.organization = branchobj.organization
            obj.branch = branchobj
            obj.amount = amount
            obj.bill = 'bill'+str(randomnum)+'.jpg'
            obj.date = todaydate
            obj.save()

            if amount <= branchobj.balance:
                branchobj.balance = branchobj.balance - amount
                branchobj.save()
            else:
                branchobj.balance = 0
                branchobj.save()

            status = 'done'

        else:
            error = 'yes'
    branchbal = branchobj.balance
    context = {
        'error': error,
        'status': status,
        'balance': branchbal
    }
    return render(request,link,context)